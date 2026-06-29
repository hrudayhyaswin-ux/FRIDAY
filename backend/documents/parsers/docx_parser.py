import os
import logging

logger = logging.getLogger(__name__)

def parse_docx(file_path: str) -> str:
    """Extracts text from a DOCX file using python-docx."""
    text = ""
    try:
        import docx
        logger.info(f"Parsing DOCX: {file_path}")
        doc = docx.Document(file_path)
        paragraphs_text = [para.text for para in doc.paragraphs]
        
        # Also parse tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                paragraphs_text.append(" | ".join(row_text))
                
        text = "\n".join(paragraphs_text)
        return text
    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_path}: {e}")
        raise e
